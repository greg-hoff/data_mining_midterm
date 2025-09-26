import pandas as pd
from mlxtend.frequent_patterns import apriori
from mlxtend.frequent_patterns import association_rules


from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches

doc = Document()
section = doc.sections[-1]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(11.0)
section.page_height = Inches(8.5)

doc.add_heading('Apriori Analysis - Transaction 1', 0)

data = pd.read_csv('trans1_onehot.csv')
pd.set_option('display.max_rows', None)  # Show all rows
pd.set_option('display.max_columns', None)  # Show all columns
data = data.astype(bool)  # Convert 0/1 to True/False
doc.add_paragraph(str(data)) # Print full DataFrame to Word

# Find frequent itemsets with a minimum support of 0.1
frequent_itemsets = apriori(data, min_support=0.1, use_colnames=True)
doc.add_paragraph(str(frequent_itemsets))

# Generate rules with a minimum confidence of 0.5
rules = association_rules(frequent_itemsets, metric="confidence", min_threshold=0.5)
doc.add_paragraph(str(rules))

doc.save('hoffer_gregory_midtermproj1.docx')